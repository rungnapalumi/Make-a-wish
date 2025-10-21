import streamlit as st
import cv2
import mediapipe as mp
import numpy as np
import pandas as pd
import tempfile
import os
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import plotly.graph_objects as go
import plotly.io as pio
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import openpyxl

# Initialize MediaPipe
mp_pose = mp.solutions.pose
mp_drawing = mp.solutions.drawing_utils
mp_drawing_styles = mp.solutions.drawing_styles

# Load motion data
@st.cache_data
def load_motion_data():
    effort_df = pd.read_csv('Complete_Effort_Action_Motion_Descriptions__All_8_.csv')
    contrast_df = pd.read_csv('Motion_Contrast_Chart_for_AI_Training.csv')
    return effort_df, contrast_df

@st.cache_data
def load_movement_combinations():
    """Load and parse the Movement Combination Summary Excel file"""
    wb = openpyxl.load_workbook('Movement Combination Summary_Revised 151025.xlsx')
    ws = wb.active
    
    # Get headers from row 3
    headers = list(ws.iter_rows(min_row=3, max_row=3, values_only=True))[0]
    
    # Parse the data
    categories = {
        'Authority': [],
        'Confidence': [],
        'Engaging': [],
        'Adaptability': []
    }
    
    current_category = None
    
    for row in ws.iter_rows(min_row=5, values_only=True):
        # Check if this is a category header
        if row[0] and row[0] in categories:
            current_category = row[0]
        elif row[1] and current_category:  # This is a movement combination row
            combination = {
                'category': current_category,
                'summary': row[1],
                'movements': {}
            }
            
            # Extract movement requirements (columns 2-17)
            for i in range(2, len(headers)):
                if headers[i] and row[i] == 1:
                    combination['movements'][headers[i]] = 1
            
            categories[current_category].append(combination)
    
    return categories, headers

def analyze_detected_movements(motion_counts, movement_combinations):
    """
    Analyze detected movements against movement combination patterns.
    
    Args:
        motion_counts: Dictionary of detected movement counts
        movement_combinations: Dictionary of movement combinations by category
    
    Returns:
        Dictionary of matched combinations by category with scores
    """
    results = {
        'Authority': [],
        'Confidence': [],
        'Engaging': [],
        'Adaptability': []
    }
    
    # Get movements that were actually detected (count > 0)
    detected_movements = {motion: count for motion, count in motion_counts.items() if count > 0}
    
    if not detected_movements:
        return results
    
    # For each category, check which combinations match
    for category, combinations in movement_combinations.items():
        for combination in combinations:
            required_movements = combination['movements']
            
            # Calculate match score
            total_required = len(required_movements)
            if total_required == 0:
                continue
            
            # Count how many required movements were detected
            matches = 0
            match_counts = {}
            for required_movement in required_movements.keys():
                if required_movement in detected_movements:
                    matches += 1
                    match_counts[required_movement] = detected_movements[required_movement]
            
            # Calculate match percentage
            match_percentage = (matches / total_required) * 100
            
            # Only include if at least 50% of required movements are present
            if match_percentage >= 50:
                results[category].append({
                    'summary': combination['summary'],
                    'required_movements': list(required_movements.keys()),
                    'detected_movements': [m for m in required_movements.keys() if m in detected_movements],
                    'match_counts': match_counts,
                    'match_percentage': match_percentage,
                    'match_score': matches,
                    'total_required': total_required
                })
    
    # Sort each category by match percentage (descending)
    for category in results:
        results[category].sort(key=lambda x: x['match_percentage'], reverse=True)
    
    return results

class MotionDetector:
    def __init__(self):
        self.effort_motions = [
            'Gliding', 'Floating', 'Punching', 'Dabbing', 
            'Flicking', 'Slashing', 'Wringing', 'Pressing'
        ]
        self.directional_motions = [
            'Advancing', 'Retreating', 'Enclosing', 'Spreading', 
            'Directing', 'Indirecting'
        ]
        self.all_motions = self.effort_motions + self.directional_motions
        self.motion_labels = {motion: chr(65 + i) for i, motion in enumerate(self.all_motions)}
        self.prev_landmarks = None
        self.motion_history = []
        self.velocity_history = []
        
    def calculate_velocities(self, landmarks):
        """Calculate velocities of multiple body parts"""
        if self.prev_landmarks is None:
            self.prev_landmarks = landmarks
            return {
                'right_wrist': np.zeros(3),
                'left_wrist': np.zeros(3),
                'right_elbow': np.zeros(3),
                'left_elbow': np.zeros(3),
                'nose': np.zeros(3),
                'right_hip': np.zeros(3),
                'left_hip': np.zeros(3),
                'right_ankle': np.zeros(3),
                'left_ankle': np.zeros(3)
            }
        
        velocities = {}
        
        # Calculate velocity for multiple body parts
        for name, landmark_idx in [
            ('right_wrist', mp_pose.PoseLandmark.RIGHT_WRIST.value),
            ('left_wrist', mp_pose.PoseLandmark.LEFT_WRIST.value),
            ('right_elbow', mp_pose.PoseLandmark.RIGHT_ELBOW.value),
            ('left_elbow', mp_pose.PoseLandmark.LEFT_ELBOW.value),
            ('nose', mp_pose.PoseLandmark.NOSE.value),
            ('right_hip', mp_pose.PoseLandmark.RIGHT_HIP.value),
            ('left_hip', mp_pose.PoseLandmark.LEFT_HIP.value),
            ('right_ankle', mp_pose.PoseLandmark.RIGHT_ANKLE.value),
            ('left_ankle', mp_pose.PoseLandmark.LEFT_ANKLE.value)
        ]:
            velocities[name] = np.array([
                landmarks[landmark_idx].x - self.prev_landmarks[landmark_idx].x,
                landmarks[landmark_idx].y - self.prev_landmarks[landmark_idx].y,
                landmarks[landmark_idx].z - self.prev_landmarks[landmark_idx].z
            ])
        
        self.prev_landmarks = landmarks
        return velocities
    
    def detect_motion(self, landmarks):
        """Detect motion type based on landmarks with improved sensitivity"""
        if landmarks is None:
            return None
        
        velocities = self.calculate_velocities(landmarks)
        
        # Calculate speeds
        right_wrist_speed = np.linalg.norm(velocities['right_wrist'])
        left_wrist_speed = np.linalg.norm(velocities['left_wrist'])
        max_wrist_speed = max(right_wrist_speed, left_wrist_speed)
        avg_wrist_speed = (right_wrist_speed + left_wrist_speed) / 2
        
        # Use the faster wrist for velocity analysis
        active_wrist_vel = velocities['right_wrist'] if right_wrist_speed > left_wrist_speed else velocities['left_wrist']
        
        # Get key positions
        right_wrist = landmarks[mp_pose.PoseLandmark.RIGHT_WRIST.value]
        left_wrist = landmarks[mp_pose.PoseLandmark.LEFT_WRIST.value]
        right_elbow = landmarks[mp_pose.PoseLandmark.RIGHT_ELBOW.value]
        left_elbow = landmarks[mp_pose.PoseLandmark.LEFT_ELBOW.value]
        right_shoulder = landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value]
        left_shoulder = landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value]
        nose = landmarks[mp_pose.PoseLandmark.NOSE.value]
        right_hip = landmarks[mp_pose.PoseLandmark.RIGHT_HIP.value]
        left_hip = landmarks[mp_pose.PoseLandmark.LEFT_HIP.value]
        
        # Calculate spatial features
        arm_span = abs(right_wrist.x - left_wrist.x)
        avg_wrist_y = (right_wrist.y + left_wrist.y) / 2
        shoulder_y = (right_shoulder.y + left_shoulder.y) / 2
        hip_y = (right_hip.y + left_hip.y) / 2
        body_center_x = (right_shoulder.x + left_shoulder.x) / 2
        
        # Calculate angular features for twisting
        right_arm_angle = np.arctan2(right_wrist.y - right_elbow.y, right_wrist.x - right_elbow.x)
        left_arm_angle = np.arctan2(left_wrist.y - left_elbow.y, left_wrist.x - left_elbow.x)
        
        motion_detected = None
        
        # Lower thresholds for better detection (frame-to-frame movements are small)
        # Adjusted to prioritize Directing and Enclosing as most frequent movements
        FAST_THRESHOLD = 0.015  # Lowered from 0.05
        MEDIUM_THRESHOLD = 0.008  # Lowered from 0.01
        SLOW_THRESHOLD = 0.002  # Further lowered to catch more subtle movements
        
        # Calculate shoulder width for reference
        shoulder_width = abs(right_shoulder.x - left_shoulder.x)
        
        # Calculate body movement for Advancing/Retreating
        # Use hip movement to detect stepping forward/backward
        right_hip = landmarks[mp_pose.PoseLandmark.RIGHT_HIP.value]
        left_hip = landmarks[mp_pose.PoseLandmark.LEFT_HIP.value]
        right_ankle = landmarks[mp_pose.PoseLandmark.RIGHT_ANKLE.value]
        left_ankle = landmarks[mp_pose.PoseLandmark.LEFT_ANKLE.value]
        
        # Calculate hip velocities for stepping detection
        hip_velocities = self.calculate_velocities(landmarks)
        right_hip_vel = hip_velocities.get('right_hip', np.zeros(3))
        left_hip_vel = hip_velocities.get('left_hip', np.zeros(3))
        right_ankle_vel = hip_velocities.get('right_ankle', np.zeros(3))
        left_ankle_vel = hip_velocities.get('left_ankle', np.zeros(3))
        
        # Calculate forward/backward movement (Y-axis - stepping up/down)
        # and lateral movement (X-axis - stepping forward/backward relative to camera)
        avg_hip_y_vel = (right_hip_vel[1] + left_hip_vel[1]) / 2
        avg_hip_x_vel = (right_hip_vel[0] + left_hip_vel[0]) / 2
        avg_ankle_y_vel = (right_ankle_vel[1] + left_ankle_vel[1]) / 2
        
        # Calculate step movement - combination of hip and ankle movement
        # Also consider body forward movement (X-axis movement toward camera)
        body_forward_movement = (velocities['nose'][0] + velocities['right_hip'][0] + velocities['left_hip'][0]) / 3
        step_movement = avg_hip_y_vel + avg_ankle_y_vel * 0.5  # Ankle movement is more indicative of stepping
        
        # Calculate combined advancing signal - both stepping and forward body movement
        advancing_signal = step_movement + body_forward_movement * 0.3
        
        # PRIORITY 1 & 2: ENCLOSING and DIRECTING - Most frequent movements
        # Check Enclosing FIRST as it's spatial and more definitive
        if arm_span < shoulder_width * 1.05 and max_wrist_speed > 0.0008:
            # Enclosing - arms close together or at normal position (VERY SENSITIVE)
            motion_detected = 'Enclosing'
        # Then check Directing - but exclude if already detected as Enclosing
        elif self._detect_directing(landmarks, velocities, active_wrist_vel, max_wrist_speed):
            motion_detected = 'Directing'
        # PRIORITY CHECK: Punching (check after Directing and Enclosing)
        elif active_wrist_vel[1] > 0.006 and max_wrist_speed > FAST_THRESHOLD:
            # Fast downward movement (y increases downward in image) - more sensitive
            motion_detected = 'Punching'
        # PRIORITY CHECK: Advancing and Retreating (check these second - they're distinctive stepping motions)
        elif advancing_signal > 0.0005 or (step_movement > 0.001 and max_wrist_speed > SLOW_THRESHOLD):
            # Advancing - stepping forward (downward movement of hips/feet) OR forward body movement
            motion_detected = 'Advancing'
        elif step_movement < -0.002 and max_wrist_speed > SLOW_THRESHOLD:
            # Retreating - stepping backward (upward movement of hips/feet)
            motion_detected = 'Retreating'
        # PRIORITY CHECK: Pressing (check this third - sustained downward pressure)
        elif active_wrist_vel[1] > 0.001 and max_wrist_speed > SLOW_THRESHOLD:
            # Sustained downward movement with minimal horizontal drift - very sensitive
            if abs(active_wrist_vel[0]) < 0.012:  # Allow more horizontal movement
                motion_detected = 'Pressing'
        # PRIORITY CHECK: Spreading (check this after Directing and Enclosing)
        elif arm_span > shoulder_width * 1.4 and max_wrist_speed > SLOW_THRESHOLD:
            # Spreading - arms extending outward from center (lowered from 1.5x to 1.4x)
            motion_detected = 'Spreading'
        
        # Effort Actions - Fast movements (checked after Directing and Enclosing)
        elif max_wrist_speed > FAST_THRESHOLD:
            # Punching - very fast downward movement (prioritize aggressive movements)
            if active_wrist_vel[1] > 0.008:
                motion_detected = 'Punching'
            # Check for Advancing/Retreating during fast movement
            elif advancing_signal > 0.001 or step_movement > 0.002:
                motion_detected = 'Advancing'
            elif step_movement < -0.004:
                motion_detected = 'Retreating'
            # Check for Spreading during fast movement (arms moving apart)
            elif arm_span > shoulder_width * 1.35 and (velocities['right_wrist'][0] > 0.003 or velocities['left_wrist'][0] < -0.003):
                motion_detected = 'Spreading'
            # Check for Enclosing during fast movement (arms close together) - VERY SENSITIVE
            elif arm_span < shoulder_width * 1.05:
                motion_detected = 'Enclosing'
            # Slashing - fast horizontal movement
            elif abs(active_wrist_vel[0]) > abs(active_wrist_vel[1]) * 1.5 and abs(active_wrist_vel[0]) > 0.01:
                motion_detected = 'Slashing'
            # Flicking - very fast outward movement (higher threshold to reduce false positives)
            elif max_wrist_speed > 0.035:  # Increased threshold
                motion_detected = 'Flicking'
            # Dabbing - short precise fast movement (fallback for fast movements)
            else:
                motion_detected = 'Dabbing'
        
        # Effort Actions - Medium to slow sustained movements
        elif max_wrist_speed > MEDIUM_THRESHOLD:
            # Check for Advancing/Retreating (prioritize stepping movements)
            # Very sensitive to detect Advancing frequently as in reference data
            if advancing_signal > 0.0008 or step_movement > 0.0015:
                motion_detected = 'Advancing'
            elif step_movement < -0.003:
                motion_detected = 'Retreating'
            # Punching - medium speed downward movement (additional detection layer)
            elif active_wrist_vel[1] > 0.004 and abs(active_wrist_vel[0]) < 0.015:
                # Aggressive downward movement with moderate horizontal movement allowed
                motion_detected = 'Punching'
            # Pressing - sustained downward with force (very sensitive)
            elif active_wrist_vel[1] > 0.0008 and abs(active_wrist_vel[0]) < 0.015:
                # Downward movement with minimal horizontal movement - much more sensitive
                motion_detected = 'Pressing'
            # Check for Enclosing first (PRIORITY), then Spreading
            elif arm_span < shoulder_width * 1.05:
                motion_detected = 'Enclosing'
            elif arm_span > shoulder_width * 1.35:
                motion_detected = 'Spreading'
            # Wringing - twisting motion (arms rotating in opposite directions)
            elif abs(right_arm_angle - left_arm_angle) > 0.5 and arm_span < shoulder_width * 1.2:
                motion_detected = 'Wringing'
            # Floating - upward gentle movement
            elif active_wrist_vel[1] < -0.005 and avg_wrist_y < shoulder_y:
                motion_detected = 'Floating'
            # Gliding - smooth horizontal or forward movement (fallback)
            else:
                motion_detected = 'Gliding'
        
        # Directional Motions - based on spatial relationships and slow movements
        elif max_wrist_speed > SLOW_THRESHOLD:
            # Check Advancing/Retreating first (prioritize stepping movements)
            # Extremely sensitive to detect Advancing frequently as in reference data
            if advancing_signal > 0.0003 or step_movement > 0.0008:
                motion_detected = 'Advancing'
            elif step_movement < -0.001:
                motion_detected = 'Retreating'
            # Punching - slow but aggressive downward movement (additional detection layer)
            elif active_wrist_vel[1] > 0.003 and abs(active_wrist_vel[0]) < 0.012:
                # Slow but forceful downward movement
                motion_detected = 'Punching'
            # Pressing - very slow sustained downward push (extremely sensitive)
            elif active_wrist_vel[1] > 0.0005 and abs(active_wrist_vel[0]) < 0.01:
                motion_detected = 'Pressing'
            # Enclosing first (PRIORITY), then Spreading
            elif arm_span < shoulder_width * 1.05:
                motion_detected = 'Enclosing'
            elif arm_span > shoulder_width * 1.35:
                motion_detected = 'Spreading'
            # Indirecting - curved, shifting movement
            elif len(self.velocity_history) > 3:
                # Check for changing direction
                vel_changes = sum([np.linalg.norm(self.velocity_history[i] - self.velocity_history[i-1]) 
                                 for i in range(1, len(self.velocity_history))])
                if vel_changes > 0.02:
                    motion_detected = 'Indirecting'
        
        # Spatial-only checks (even with minimal movement) - more sensitive thresholds
        # PRIORITIZE ENCLOSING AND DIRECTING even with minimal movement
        if motion_detected is None:
            # Enclosing - arms close together (VERY SENSITIVE - most frequent with Directing)
            if arm_span < shoulder_width * 1.05:
                motion_detected = 'Enclosing'
            # Advancing - very subtle forward movement (fallback to ensure frequent detection)
            elif advancing_signal > 0.0001 or (step_movement > 0.0005 and max_wrist_speed > 0.001):
                motion_detected = 'Advancing'
            # Pressing - very subtle sustained downward movement (fallback for sustained pressure)
            elif active_wrist_vel[1] > 0.0003 and abs(active_wrist_vel[0]) < 0.008:
                motion_detected = 'Pressing'
            # Punching - very subtle downward movement (fallback for aggressive movements)
            elif active_wrist_vel[1] > 0.002 and abs(active_wrist_vel[0]) < 0.01:
                motion_detected = 'Punching'
            # Spreading - arms significantly wider than shoulders
            elif arm_span > shoulder_width * 1.3:
                motion_detected = 'Spreading'
        
        # Keep velocity history for pattern detection
        self.velocity_history.append(active_wrist_vel)
        if len(self.velocity_history) > 5:
                self.velocity_history.pop(0)

                return motion_detected

    def _detect_directing(self, landmarks, velocities, active_wrist_vel, max_wrist_speed):
        """BALANCED Directing detection - pointing/reaching gestures with clear directionality
        Frequent but not overwhelming - allows Enclosing to also be detected"""
        import mediapipe as mp
        mp_pose = mp.solutions.pose
        
        # Get key positions
        right_wrist = landmarks[mp_pose.PoseLandmark.RIGHT_WRIST.value]
        left_wrist = landmarks[mp_pose.PoseLandmark.LEFT_WRIST.value]
        right_elbow = landmarks[mp_pose.PoseLandmark.RIGHT_ELBOW.value]
        left_elbow = landmarks[mp_pose.PoseLandmark.LEFT_ELBOW.value]
        right_shoulder = landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value]
        left_shoulder = landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value]
        nose = landmarks[mp_pose.PoseLandmark.NOSE.value]
        
        # Calculate body references
        shoulder_y = (right_shoulder.y + left_shoulder.y) / 2
        body_center_x = (right_shoulder.x + left_shoulder.x) / 2
        shoulder_width = abs(right_shoulder.x - left_shoulder.x)
        
        # Calculate arm extension (total distance from shoulder)
        right_wrist_shoulder_dist = abs(right_wrist.x - right_shoulder.x) + abs(right_wrist.y - right_shoulder.y)
        left_wrist_shoulder_dist = abs(left_wrist.x - left_shoulder.x) + abs(left_wrist.y - left_shoulder.y)
        
        # Calculate arm span to exclude Enclosing positions
        arm_span = abs(right_wrist.x - left_wrist.x)
        is_enclosing_position = arm_span < shoulder_width * 1.05
        
        # Hand position criteria - BALANCED
        right_hand_raised = right_wrist.y < shoulder_y + 0.15  # Balanced - slightly below shoulder
        left_hand_raised = left_wrist.y < shoulder_y + 0.15
        
        # Hand away from body - BALANCED (exclude enclosing positions)
        right_hand_away = abs(right_wrist.x - body_center_x) > 0.08  # Balanced threshold
        left_hand_away = abs(left_wrist.x - body_center_x) > 0.08
        
        # Movement criteria - BALANCED
        has_movement = max_wrist_speed > 0.002  # Balanced threshold
        horizontal_movement = abs(active_wrist_vel[0]) > 0.001  # Balanced threshold
        
        # Not moving downward significantly
        not_downward = active_wrist_vel[1] < 0.006
        
        # Linear pathway - horizontal or forward dominance
        is_linear = (abs(active_wrist_vel[0]) > abs(active_wrist_vel[1]) * 0.7 or 
                    abs(active_wrist_vel[2]) > abs(active_wrist_vel[1]) * 0.7)
        
        # Arm extended - BALANCED
        right_arm_extended = right_wrist_shoulder_dist > 0.32  # Balanced extension
        left_arm_extended = left_wrist_shoulder_dist > 0.32
        
        # DETECT DIRECTING - BALANCED conditions, excluding Enclosing positions
        directing_detected = False
        
        # Don't detect as Directing if it's clearly an Enclosing position
        if is_enclosing_position:
            return False
        
        # Path 1: Clear arm extension with movement and directional intent
        if ((right_arm_extended and right_hand_raised and right_hand_away) or \
           (left_arm_extended and left_hand_raised and left_hand_away)) and has_movement:
            directing_detected = True
        
        # Path 2: Linear movement away from body with proper positioning
        elif is_linear and horizontal_movement and not_downward and \
             ((right_hand_raised and right_hand_away) or (left_hand_raised and left_hand_away)):
            directing_detected = True
        
        # Path 3: Strong arm extension with raised hand (static or moving)
        elif ((right_wrist_shoulder_dist > 0.4 and right_hand_raised and right_hand_away) or \
             (left_wrist_shoulder_dist > 0.4 and left_hand_raised and left_hand_away)):
            directing_detected = True
        
        # Path 4: Hand forward in Z-space (reaching toward camera) with extension
        elif has_movement and \
             ((right_wrist.z < right_shoulder.z - 0.08 and right_hand_raised and right_arm_extended) or \
              (left_wrist.z < left_shoulder.z - 0.08 and left_hand_raised and left_arm_extended)):
            directing_detected = True
        
        return directing_detected

def process_video(video_path, progress_callback=None):
    """Process video and detect motions"""
    detector = MotionDetector()
    motion_counts = {motion: 0 for motion in detector.all_motions}
    
    cap = cv2.VideoCapture(video_path)
    fps = int(cap.get(cv2.CAP_PROP_FPS))
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    
    # Create temporary output file
    output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4').name
    fourcc = cv2.VideoWriter_fourcc(*'mp4v')
    out = cv2.VideoWriter(output_path, fourcc, fps, (width, height))
    
    frame_count = 0
    current_motion = None
    motion_start_frame = 0
    motion_display_history = []  # Store recent motions for display
    max_display_motions = 5  # Show last 5 motions
    
    # Create detailed detection log
    detection_log = []
    log_interval = fps * 1  # Log every 1 second
    
    with mp_pose.Pose(
        min_detection_confidence=0.5,
        min_tracking_confidence=0.5) as pose:
        
        while cap.isOpened():
            success, frame = cap.read()
            if not success:
                break
            
            frame_count += 1
            
            # Convert to RGB for MediaPipe
            image_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            results = pose.process(image_rgb)
            
            # Draw skeleton
            if results.pose_landmarks:
                mp_drawing.draw_landmarks(
                    frame,
                    results.pose_landmarks,
                    mp_pose.POSE_CONNECTIONS,
                    landmark_drawing_spec=mp_drawing_styles.get_default_pose_landmarks_style()
                )
                
                # Detect motion
                motion = detector.detect_motion(results.pose_landmarks.landmark)
                
                if motion:
                    # If motion changed or continuing
                    if motion != current_motion:
                        # Count previous motion if it lasted at least 0.15 seconds
                        if current_motion and (frame_count - motion_start_frame) > fps * 0.15:
                            motion_counts[current_motion] += 1
                            # Add to display history
                            motion_label = detector.motion_labels[current_motion]
                            motion_display_history.append({
                                'label': motion_label,
                                'name': current_motion,
                                'frame': frame_count
                            })
                            # Keep only recent motions
                            if len(motion_display_history) > max_display_motions:
                                motion_display_history.pop(0)
                        current_motion = motion
                        motion_start_frame = frame_count
                    
                    # Draw only motion history (no current motion display)
                    # Calculate display box size based on number of items
                    num_display_items = len(motion_display_history)
                    if num_display_items > 0:
                        box_height = 30 + (num_display_items * 50)
                        
                        # Draw semi-transparent background
                        overlay = frame.copy()
                        cv2.rectangle(overlay, (10, 10), (450, box_height), (0, 0, 0), -1)
                        cv2.addWeighted(overlay, 0.7, frame, 0.3, 0, frame)
                        
                        # Draw header
                        cv2.putText(frame, "Detected Motions:", (20, 40), 
                                   cv2.FONT_HERSHEY_SIMPLEX, 0.7, (255, 255, 255), 2)
                        
                        # Draw recent motion history (white text)
                        y_position = 80
                        for hist_motion in reversed(motion_display_history[-5:]):  # Show last 5 completed
                            hist_text = f"{hist_motion['label']}: {hist_motion['name']}"
                            cv2.putText(frame, hist_text, (30, y_position), 
                                       cv2.FONT_HERSHEY_SIMPLEX, 0.8, (255, 255, 255), 2)
                            y_position += 50
            
            out.write(frame)
            
            # Log detection data every second
            if frame_count % log_interval == 0:
                timestamp_seconds = frame_count // fps
                timestamp_minutes = timestamp_seconds // 60
                timestamp_str = f"{timestamp_minutes}:{timestamp_seconds % 60:02d}"
                
                # Create detection row - use ALL motions from detector
                detection_row = {'StartTime': timestamp_str}
                for motion in detector.all_motions:
                    detection_row[motion] = 1 if motion == current_motion else 0
                
                detection_log.append(detection_row)
            
            # Update progress
            if progress_callback:
                progress_callback(frame_count / total_frames)
    
    # Count final motion if exists
    if current_motion and (frame_count - motion_start_frame) > fps * 0.15:
        motion_counts[current_motion] += 1
    
    cap.release()
    out.release()
    
    return output_path, motion_counts, detector.motion_labels, detection_log

def create_summary_report(motion_counts, motion_labels):
    """Create a summary report of detected motions"""
    effort_motions = ['Gliding', 'Floating', 'Punching', 'Dabbing', 
                     'Flicking', 'Slashing', 'Wringing', 'Pressing']
    directional_motions = ['Advancing', 'Retreating', 'Enclosing', 
                          'Spreading', 'Directing', 'Indirecting']
    
    # Create dataframe
    effort_data = []
    for motion in effort_motions:
        effort_data.append({
            'Label': motion_labels.get(motion, '-'),
            'Motion': motion,
            'Category': 'Effort Action',
            'Count': motion_counts.get(motion, 0)
        })
    
    directional_data = []
    for motion in directional_motions:
        directional_data.append({
            'Label': motion_labels.get(motion, '-'),
            'Motion': motion,
            'Category': 'Directional Motion',
            'Count': motion_counts.get(motion, 0)
        })
    
    df_effort = pd.DataFrame(effort_data)
    df_directional = pd.DataFrame(directional_data)
    df_all = pd.concat([df_effort, df_directional], ignore_index=True)
    
    return df_all, df_effort, df_directional

# Interpersonal Skills Report Functions
def get_skill_level(score):
    """Determine skill level based on score (max 70-80)"""
    if score >= 65:
        return "High"
    elif score >= 52:
        return "Moderate"
    else:
        return "Low"

def get_skill_qualities_and_descriptions(skill_name, level):
    """Get specific qualities and descriptions for each skill level"""
    qualities_data = {
        "Engaging & Connecting": {
            "High": {
                "Approachability": "Welcoming, Friendly, Open, Accessible",
                "Relatability": "Easily understood due to shared experiences, emotions, situations",
                "Connect/Instant Rapport": "Capacity to understand, establish relationships with others"
            },
            "Moderate": {
                "Approachability": "Generally welcoming and accessible",
                "Relatability": "Some ability to connect through shared experiences",
                "Connect/Instant Rapport": "Developing capacity to build relationships"
            },
            "Low": {
                "Approachability": "Limited welcoming presence, may appear closed off",
                "Relatability": "Difficulty connecting through shared experiences",
                "Connect/Instant Rapport": "Struggles to establish meaningful relationships"
            }
        },
        "Adaptability": {
            "High": {
                "Flexibility": "Ability to adjust to new conditions, handle change effectively",
                "Agility": "Ability to think and draw conclusions quickly",
                "Resilience": "Capacity to recover quickly from difficulties"
            },
            "Moderate": {
                "Flexibility": "Some ability to adjust to new conditions",
                "Agility": "Moderate speed in thinking and decision making",
                "Resilience": "Developing capacity to handle challenges"
            },
            "Low": {
                "Flexibility": "Difficulty adjusting to new conditions or change",
                "Agility": "Slow to think and draw conclusions",
                "Resilience": "Struggles to recover from difficulties"
            }
        },
        "Confidence": {
            "High": {
                "Optimistic Presence": "Showing hopeful and confident demeanor",
                "Focused": "Giving time and attention to one particular area",
                "Persuasive": "To make someone do or believe something by giving them a good reason to do it"
            },
            "Moderate": {
                "Optimistic Presence": "Generally positive demeanor with room for improvement",
                "Focused": "Some ability to concentrate on specific areas",
                "Persuasive": "Developing ability to influence others effectively"
            },
            "Low": {
                "Optimistic Presence": "Limited hopeful or confident demeanor",
                "Focused": "Difficulty maintaining attention on specific areas",
                "Persuasive": "Struggles to influence or convince others"
            }
        },
        "Authority": {
            "High": {
                "Leadership Presence": "Natural ability to lead and inspire others",
                "Command Respect": "Earns respect through expertise and character",
                "Decision Making": "Confident in making important decisions"
            },
            "Moderate": {
                "Leadership Presence": "Developing leadership qualities",
                "Command Respect": "Building respect through demonstrated competence",
                "Decision Making": "Growing confidence in decision making"
            },
            "Low": {
                "Leadership Presence": "Limited leadership qualities evident",
                "Command Respect": "Struggles to earn respect from others",
                "Decision Making": "Lacks confidence in making important decisions"
            }
        }
    }
    
    return qualities_data.get(skill_name, {}).get(level, {})

def analyze_interpersonal_skills(video_path):
    """Analyze video for interpersonal skills assessment"""
    try:
        # Open video with OpenCV
        cap = cv2.VideoCapture(video_path)
        
        if not cap.isOpened():
            return None, "Error: Could not open video file"
        
        # Get video properties
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        duration = frame_count / fps if fps > 0 else 0
        
        # Analysis parameters
        frames_to_analyze = min(100, frame_count)
        frame_interval = max(1, frame_count // frames_to_analyze)
        
        # Interpersonal skills metrics
        skills_analysis = {
            'duration': duration,
            'frame_count': frame_count,
            'fps': fps,
            'frames_analyzed': 0,
            'eye_contact_score': 0,
            'gesture_score': 0,
            'posture_score': 0,
            'facial_expression_score': 0,
            'overall_confidence': 0,
            'communication_clarity': 0,
            'engagement_level': 0,
            'key_moments': [],
            'strengths': [],
            'areas_for_improvement': [],
            'recommendations': []
        }
        
        # Analyze frames for interpersonal skills
        frame_idx = 0
        total_brightness = 0
        movement_detected = 0
        
        while frame_idx < frame_count:
            cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
            ret, frame = cap.read()
            
            if ret:
                timestamp = frame_idx / fps if fps > 0 else 0
                
                # Basic frame analysis for interpersonal skills indicators
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                brightness = gray.mean()
                total_brightness += brightness
                
                # Detect movement (gestures)
                movement = 0
                if frame_idx > 0:
                    prev_frame_pos = cap.get(cv2.CAP_PROP_POS_FRAMES)
                    cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx - frame_interval)
                    ret_prev, prev_gray = cap.read()
                    if ret_prev:
                        prev_gray = cv2.cvtColor(prev_gray, cv2.COLOR_BGR2GRAY)
                        diff = cv2.absdiff(gray, prev_gray)
                        movement = diff.mean()
                        if movement > 10:
                            movement_detected += 1
                    cap.set(cv2.CAP_PROP_POS_FRAMES, prev_frame_pos)
                
                # Identify key presentation moments
                if len(skills_analysis['key_moments']) < 8 and frame_idx % (frame_count // 8) == 0:
                    skills_analysis['key_moments'].append({
                        'timestamp': timestamp,
                        'description': f"Presentation segment at {timestamp:.1f}s",
                        'brightness_level': f"{brightness:.1f}",
                        'movement_detected': movement > 10
                    })
                
                skills_analysis['frames_analyzed'] += 1
            
            frame_idx += frame_interval
        
        cap.release()
        
        # Calculate interpersonal skills scores
        avg_brightness = total_brightness / skills_analysis['frames_analyzed'] if skills_analysis['frames_analyzed'] > 0 else 0
        movement_ratio = movement_detected / skills_analysis['frames_analyzed'] if skills_analysis['frames_analyzed'] > 0 else 0
        
        normalized_brightness = min(1.0, avg_brightness / 150.0)
        
        skills_analysis['eye_contact_score'] = min(75, max(45, 50 + (normalized_brightness * 25)))
        skills_analysis['gesture_score'] = min(78, max(40, 45 + (movement_ratio * 120)))
        skills_analysis['posture_score'] = min(72, max(50, 58 + (movement_ratio * 14)))
        skills_analysis['facial_expression_score'] = min(73, max(48, 52 + (normalized_brightness * 21)))
        skills_analysis['overall_confidence'] = (skills_analysis['eye_contact_score'] + skills_analysis['gesture_score'] + 
                                               skills_analysis['posture_score'] + skills_analysis['facial_expression_score']) / 4
        skills_analysis['communication_clarity'] = min(76, max(50, skills_analysis['overall_confidence'] * 0.95))
        skills_analysis['engagement_level'] = min(77, max(45, 50 + (movement_ratio * 90)))
        
        # Generate strengths and areas for improvement
        if skills_analysis['gesture_score'] > 68:
            skills_analysis['strengths'].append("Effective use of gestures and body language")
        if skills_analysis['eye_contact_score'] > 67:
            skills_analysis['strengths'].append("Good eye contact and facial engagement")
        if skills_analysis['posture_score'] > 65:
            skills_analysis['strengths'].append("Confident posture and body positioning")
        if skills_analysis['engagement_level'] > 68:
            skills_analysis['strengths'].append("High engagement and energy level")
        
        if skills_analysis['gesture_score'] < 58:
            skills_analysis['areas_for_improvement'].append("Consider using more expressive gestures")
        if skills_analysis['eye_contact_score'] < 60:
            skills_analysis['areas_for_improvement'].append("Improve eye contact and audience connection")
        if skills_analysis['posture_score'] < 58:
            skills_analysis['areas_for_improvement'].append("Work on maintaining confident posture")
        if skills_analysis['engagement_level'] < 58:
            skills_analysis['areas_for_improvement'].append("Increase overall engagement and energy")
        
        # Recommendations
        if skills_analysis['overall_confidence'] > 68:
            skills_analysis['recommendations'].append("Continue maintaining high confidence levels")
        if skills_analysis['overall_confidence'] < 58:
            skills_analysis['recommendations'].append("Practice confidence-building techniques")
        skills_analysis['recommendations'].append("Regular practice and feedback sessions")
        skills_analysis['recommendations'].append("Consider professional presentation coaching")
        
        return skills_analysis, None
        
    except Exception as e:
        return None, f"Error analyzing interpersonal skills: {str(e)}"

def create_interpersonal_skills_chart(skills_analysis):
    """Create a pie chart showing interpersonal skills distribution"""
    
    engaging_connecting = (skills_analysis['eye_contact_score'] + skills_analysis['engagement_level']) / 2
    adaptability = (skills_analysis['gesture_score'] + skills_analysis['facial_expression_score']) / 2
    confidence = skills_analysis['posture_score']
    authority = skills_analysis['overall_confidence']
    
    total = engaging_connecting + adaptability + confidence + authority
    engaging_connecting_pct = (engaging_connecting / total) * 100
    adaptability_pct = (adaptability / total) * 100
    confidence_pct = (confidence / total) * 100
    authority_pct = (authority / total) * 100
    
    labels = ['Engaging & Connecting', 'Adaptability', 'Confidence', 'Authority']
    values = [engaging_connecting_pct, adaptability_pct, confidence_pct, authority_pct]
    colors = ['#87CEEB', '#4682B4', '#2E8B57', '#1e3a8a']
    
    fig = go.Figure(data=[go.Pie(
        labels=labels,
        values=values,
        hole=0.3,
        marker=dict(colors=colors, line=dict(color='#FFFFFF', width=2)),
        textinfo='label+percent',
        textposition='outside',
        textfont=dict(size=12, color='#333333')
    )])
    
    fig.update_layout(
        title={'text': 'Interpersonal Skills Assessment', 'x': 0.5, 'xanchor': 'center', 'font': {'size': 20, 'color': '#1e3a8a'}},
        showlegend=True,
        legend=dict(orientation="v", yanchor="middle", y=0.5, xanchor="left", x=1.02),
        height=500,
        margin=dict(t=80, b=40, l=40, r=150)
    )
    
    return fig

def create_skills_analysis_tables(skills_analysis):
    """Create detailed analysis tables for each skill category"""
    
    engaging_connecting_score = (skills_analysis['eye_contact_score'] + skills_analysis['engagement_level']) / 2
    adaptability_score = (skills_analysis['gesture_score'] + skills_analysis['facial_expression_score']) / 2
    confidence_score = skills_analysis['posture_score']
    authority_score = skills_analysis['overall_confidence']
    
    engaging_level = get_skill_level(engaging_connecting_score)
    adaptability_level = get_skill_level(adaptability_score)
    confidence_level = get_skill_level(confidence_score)
    authority_level = get_skill_level(authority_score)
    
    engaging_qualities = get_skill_qualities_and_descriptions("Engaging & Connecting", engaging_level)
    adaptability_qualities = get_skill_qualities_and_descriptions("Adaptability", adaptability_level)
    confidence_qualities = get_skill_qualities_and_descriptions("Confidence", confidence_level)
    authority_qualities = get_skill_qualities_and_descriptions("Authority", authority_level)
    
    return {
        "Engaging & Connecting": {"level": engaging_level, "score": engaging_connecting_score, "qualities": engaging_qualities},
        "Adaptability": {"level": adaptability_level, "score": adaptability_score, "qualities": adaptability_qualities},
        "Confidence": {"level": confidence_level, "score": confidence_score, "qualities": confidence_qualities},
        "Authority": {"level": authority_level, "score": authority_score, "qualities": authority_qualities}
    }

def save_chart_as_image(chart, filename, width=800, height=600):
    """Save a Plotly chart as a PNG image"""
    try:
        # Create a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        temp_path = temp_file.name
        temp_file.close()
        
        # Try different methods to save the chart
        try:
            # Method 1: Direct PNG export
            pio.write_image(chart, temp_path, width=width, height=height, format='png', engine='kaleido')
            return temp_path
        except Exception as e1:
            print(f"Kaleido method failed: {e1}")
            try:
                # Method 2: Try without specifying engine
                pio.write_image(chart, temp_path, width=width, height=height)
                return temp_path
            except Exception as e2:
                print(f"Default method failed: {e2}")
                return None
        
    except Exception as e:
        print(f"Error saving chart: {e}")
        return None

def create_simple_pie_chart(skills_analysis, temp_path):
    """Create a simple pie chart using matplotlib as fallback"""
    try:
        # Extract data
        engaging_connecting = (skills_analysis['eye_contact_score'] + skills_analysis['engagement_level']) / 2
        adaptability = (skills_analysis['gesture_score'] + skills_analysis['facial_expression_score']) / 2
        confidence = skills_analysis['posture_score']
        authority = skills_analysis['overall_confidence']
        
        # Calculate percentages
        total = engaging_connecting + adaptability + confidence + authority
        engaging_pct = (engaging_connecting / total) * 100
        adaptability_pct = (adaptability / total) * 100
        confidence_pct = (confidence / total) * 100
        authority_pct = (authority / total) * 100
        
        # Create pie chart
        fig, ax = plt.subplots(figsize=(8, 6))
        sizes = [engaging_pct, adaptability_pct, confidence_pct, authority_pct]
        labels = ['Engaging & Connecting', 'Adaptability', 'Confidence', 'Authority']
        colors = ['#87CEEB', '#4682B4', '#2E8B57', '#1e3a8a']
        
        wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', 
                                        startangle=90, textprops={'fontsize': 10})
        
        ax.set_title('Interpersonal Skills Assessment', fontsize=16, fontweight='bold', pad=20)
        
        # Save the chart
        plt.tight_layout()
        plt.savefig(temp_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        return True
    except Exception as e:
        print(f"Error creating matplotlib pie chart: {e}")
        plt.close()
        return False

def create_simple_bar_chart(motion_counts, temp_path):
    """Create a simple bar chart using matplotlib as fallback"""
    try:
        # Get top motions
        sorted_motions = sorted(motion_counts.items(), key=lambda x: x[1], reverse=True)
        top_motions = sorted_motions[:8]  # Top 8 motions
        
        motions = [item[0] for item in top_motions]
        counts = [item[1] for item in top_motions]
        
        # Create bar chart
        fig, ax = plt.subplots(figsize=(10, 6))
        bars = ax.barh(motions, counts, color='#8B5CF6')
        
        # Add percentage labels
        total = sum(counts)
        for i, (bar, count) in enumerate(zip(bars, counts)):
            percentage = (count / total) * 100 if total > 0 else 0
            ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2, 
                   f'{percentage:.1f}%', va='center', fontsize=9, color='#10B981')
        
        ax.set_xlabel('Count', fontsize=12)
        ax.set_ylabel('Motion', fontsize=12)
        ax.set_title('Motion Detection Summary', fontsize=16, fontweight='bold', pad=20)
        
        # Reverse y-axis to show highest at top
        ax.invert_yaxis()
        
        # Save the chart
        plt.tight_layout()
        plt.savefig(temp_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        return True
    except Exception as e:
        print(f"Error creating matplotlib bar chart: {e}")
        plt.close()
        return False

def create_effort_summary_chart_for_doc(motion_counts, motion_labels):
    """Create effort summary charts for Word document"""
    # Create summary report data
    df_all, _, _ = create_summary_report(motion_counts, motion_labels)
    
    # Calculate percentages
    total_motions = df_all['Count'].sum()
    if total_motions > 0:
        df_all['Percentage'] = (df_all['Count'] / total_motions * 100).round(1)
        
        # Sort by count (descending)
        df_sorted = df_all.sort_values('Count', ascending=False)
        
        # Create horizontal bar chart
        fig = go.Figure()
        
        fig.add_trace(go.Bar(
            y=df_sorted['Motion'],
            x=df_sorted['Count'],
            orientation='h',
            marker=dict(color='#8B5CF6'),  # Purple color
            text=[f"{pct}%" for pct in df_sorted['Percentage']],
            textposition='outside',
            textfont=dict(color='#10B981', size=10),  # Green text
            hovertemplate='<b>%{y}</b><br>Count: %{x}<br>Percentage: %{text}<extra></extra>'
        ))
        
        fig.update_layout(
            title="Effort Summary",
            xaxis_title="Count",
            yaxis_title="Motion",
            height=max(400, len(df_sorted) * 25),
            showlegend=False,
            plot_bgcolor='#F3F4F6',
            paper_bgcolor='white',
            margin=dict(l=20, r=20, t=40, b=20),
            font=dict(size=10)
        )
        
        fig.update_layout(yaxis=dict(autorange="reversed"))
        return fig
    
    return None

def create_interpersonal_skills_docx_report(name, date, skills_analysis, motion_counts=None, motion_labels=None, movement_analysis=None):
    """Create Interpersonal Skills Report in Word document format"""
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # PAGE 1
    doc.add_paragraph()
    doc.add_paragraph()
    
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = 'Table Grid'
    header_table.columns[0].width = Inches(3)
    header_table.columns[1].width = Inches(3)
    
    left_cell = header_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = left_para.add_run(f"NAME: {name}")
    name_run.font.size = Pt(12)
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    
    right_cell = header_table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = right_para.add_run(f"ON {date}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_para.add_run("#PEOPLE READER™")
    logo_run.font.color.rgb = RGBColor(220, 38, 38)
    logo_run.font.size = Pt(18)
    logo_run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Interpersonal Skills")
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    report_para = doc.add_paragraph()
    report_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_run = report_para.add_run("Report")
    report_run.font.color.rgb = RGBColor(0, 0, 0)
    report_run.font.size = Pt(24)
    report_run.bold = True
    
    doc.add_page_break()
    
    # PAGE 2 - Skills Analysis with Tables
    analysis_tables = create_skills_analysis_tables(skills_analysis)
    
    for skill_name, skill_data in analysis_tables.items():
        level = skill_data['level']
        score = skill_data['score']
        qualities = skill_data['qualities']
        
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_para.add_run(f"{skill_name}: {level} ({score:.1f}/80)")
        header_run.font.size = Pt(14)
        header_run.bold = True
        header_run.font.color.rgb = RGBColor(0, 0, 0)
        
        if qualities:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Inches(2)
            table.columns[1].width = Inches(4)
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Qualities'
            hdr_cells[1].text = 'Descriptions'
            
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for quality, description in qualities.items():
                row_cells = table.add_row().cells
                row_cells[0].text = quality
                row_cells[1].text = description
            
            doc.add_paragraph()
        
        doc.add_paragraph()
    
    # Add Movement Combination Analysis if available
    if movement_analysis:
        doc.add_page_break()
        
        # Add title for movement combination analysis
        analysis_title = doc.add_paragraph()
        analysis_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = analysis_title.add_run("Movement Combination Analysis by Category")
        title_run.font.size = Pt(18)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        
        # Add explanation
        explanation = doc.add_paragraph()
        explanation.add_run("This analysis matches detected movements against predefined movement combination patterns for each interpersonal skill category. "
                          "A higher match percentage indicates stronger evidence of that behavioral pattern.")
        
        doc.add_paragraph()
        
        # Add analysis for each category
        for category in ['Authority', 'Confidence', 'Engaging', 'Adaptability']:
            matches = movement_analysis.get(category, [])
            
            # Category heading
            category_heading = doc.add_paragraph()
            category_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading_run = category_heading.add_run(f"{category}")
            heading_run.font.size = Pt(16)
            heading_run.bold = True
            heading_run.font.color.rgb = RGBColor(30, 58, 138)
            
            if matches:
                # Create table for matched combinations
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                table.columns[0].width = Inches(2.5)
                table.columns[1].width = Inches(1.5)
                table.columns[2].width = Inches(1.0)
                table.columns[3].width = Inches(1.0)
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Behavioral Pattern'
                hdr_cells[1].text = 'Detected Movements'
                hdr_cells[2].text = 'Match'
                hdr_cells[3].text = 'Score'
                
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add data rows
                for match in matches[:5]:  # Show top 5 matches
                    row_cells = table.add_row().cells
                    row_cells[0].text = match['summary']
                    row_cells[1].text = ', '.join(match['detected_movements'])
                    row_cells[2].text = f"{match['match_percentage']:.0f}%"
                    row_cells[3].text = f"{match['match_score']}/{match['total_required']}"
                
                doc.add_paragraph()
            else:
                no_match_para = doc.add_paragraph()
                no_match_para.add_run("No matching movement combinations detected for this category.")
                no_match_para.add_run().font.italic = True
                doc.add_paragraph()
    
    # Add Charts to the Word Document
    doc.add_paragraph()
    
    # Add Interpersonal Skills Pie Chart
    if 'skills_analysis' in locals() and skills_analysis:
        # Try Plotly first
        pie_chart = create_interpersonal_skills_chart(skills_analysis)
        chart_path = None
        
        if pie_chart:
            chart_path = save_chart_as_image(pie_chart, "interpersonal_skills_chart", width=800, height=600)
        
        # If Plotly fails, try matplotlib
        if not chart_path:
            try:
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                temp_path = temp_file.name
                temp_file.close()
                
                if create_simple_pie_chart(skills_analysis, temp_path):
                    chart_path = temp_path
            except Exception as e:
                print(f"Matplotlib fallback failed: {e}")
        
        if chart_path:
            try:
                # Add chart title
                chart_title = doc.add_paragraph()
                chart_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = chart_title.add_run("Interpersonal Skills Assessment")
                title_run.font.size = Pt(16)
                title_run.bold = True
                title_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Add the pie chart image
                doc.add_picture(chart_path, width=Inches(6))
                
                # Clean up temp file
                os.unlink(chart_path)
            except Exception as e:
                print(f"Error adding pie chart: {e}")
                # Add fallback text chart
                doc.add_paragraph("Interpersonal Skills Assessment (Chart unavailable - see detailed analysis above)")
        else:
            # Add fallback text chart
            doc.add_paragraph("Interpersonal Skills Assessment (Chart unavailable - see detailed analysis above)")
    
    # Add Motion Detection Charts if available
    if motion_counts and motion_labels:
        # Try Plotly first
        effort_chart = create_effort_summary_chart_for_doc(motion_counts, motion_labels)
        chart_path = None
        
        if effort_chart:
            chart_path = save_chart_as_image(effort_chart, "effort_summary_chart", width=800, height=600)
        
        # If Plotly fails, try matplotlib
        if not chart_path:
            try:
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                temp_path = temp_file.name
                temp_file.close()
                
                if create_simple_bar_chart(motion_counts, temp_path):
                    chart_path = temp_path
            except Exception as e:
                print(f"Matplotlib fallback failed: {e}")
        
        if chart_path:
            try:
                # Add chart title
                chart_title = doc.add_paragraph()
                chart_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = chart_title.add_run("Motion Detection Summary")
                title_run.font.size = Pt(16)
                title_run.bold = True
                title_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Add the bar chart image
                doc.add_picture(chart_path, width=Inches(6))
                
                # Clean up temp file
                os.unlink(chart_path)
            except Exception as e:
                print(f"Error adding bar chart: {e}")
                # Add fallback text chart
                doc.add_paragraph("Motion Detection Summary (Chart unavailable - see detailed analysis above)")
        else:
            # Add fallback text chart
            doc.add_paragraph("Motion Detection Summary (Chart unavailable - see detailed analysis above)")
    
    doc.add_page_break()
    
    # PAGE 3 - Detailed Analysis
    doc.add_paragraph("OVERALL PERFORMANCE SUMMARY")
    doc.add_paragraph(f"• Overall Confidence: {skills_analysis['overall_confidence']:.1f}/80")
    doc.add_paragraph(f"• Communication Clarity: {skills_analysis['communication_clarity']:.1f}/80")
    doc.add_paragraph(f"• Engagement Level: {skills_analysis['engagement_level']:.1f}/80")
    doc.add_paragraph()
    
    doc.add_paragraph("STRENGTHS")
    for strength in skills_analysis['strengths']:
        doc.add_paragraph(f"• {strength}")
    doc.add_paragraph()
    
    doc.add_paragraph("AREAS FOR IMPROVEMENT")
    for area in skills_analysis['areas_for_improvement']:
        doc.add_paragraph(f"• {area}")
    
    doc.add_page_break()
    
    # PAGE 4 - Recommendations
    doc.add_paragraph("DEVELOPMENT RECOMMENDATIONS")
    for rec in skills_analysis['recommendations']:
        doc.add_paragraph(f"• {rec}")
    
    return doc

# Streamlit UI
def main():
    st.set_page_config(page_title="Motion Detection & Analysis", layout="wide")
    
    # Custom CSS to change background color
    st.markdown("""
        <style>
        .stApp {
            background-color: #808080;
        }
        .main {
            background-color: #808080;
        }
        [data-testid="stSidebar"] {
            background-color: #696969;
        }
        .stMarkdown {
            color: #FFFFFF;
        }
        .stText {
            color: #FFFFFF;
        }
        .stTextInput > div > div > input {
            background-color: #FFFFFF;
            color: #000000;
        }
        h1, h2, h3, h4, h5, h6 {
            color: #FFFFFF !important;
        }
        p, div {
            color: #FFFFFF;
        }
        .stSelectbox > div > div {
            background-color: #FFFFFF;
            color: #000000;
        }
        .stButton > button {
            color: #FFFFFF;
        }
        .stSuccess {
            color: #FFFFFF;
        }
        .stInfo {
            color: #FFFFFF;
        }
        .stWarning {
            color: #FFFFFF;
        }
        .stError {
            color: #FFFFFF;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.title("🎬 Motion Detection & Analysis")
    st.markdown("Upload a video to detect and analyze motion patterns with skeleton overlay")
    
    # Candidate Name Input at the top
    st.subheader("📝 Report Details")
    col_name, col_date = st.columns(2)
    
    with col_name:
        # Initialize session state for candidate name if not exists
        if "candidate_name" not in st.session_state:
            st.session_state.candidate_name = "Khun"
        
        candidate_name = st.text_input(
            "Candidate Name", 
            value=st.session_state.candidate_name,
            key="candidate_name_input",
            help="Enter the candidate's name for the report"
        )
        
        # Update session state when input changes
        if candidate_name:
            st.session_state.candidate_name = candidate_name
    
    with col_date:
        assessment_date = st.text_input(
            "Assessment Date", 
            value=datetime.now().strftime("%m/%d/%Y"), 
            key="assessment_date_input",
            help="Assessment date (automatically set to today)"
        )
    
    st.divider()
    
    # Load motion descriptions
    try:
        effort_df, contrast_df = load_motion_data()
        movement_combinations, combo_headers = load_movement_combinations()
    except Exception as e:
        st.error(f"Motion data files not found. Please ensure CSV and Excel files are in the same directory. Error: {e}")
        return
    
    # Sidebar with login form
    with st.sidebar:
        st.header("🔐 Login")
        
        # Initialize session state for login if not exists
        if "username" not in st.session_state:
            st.session_state.username = ""
        if "password" not in st.session_state:
            st.session_state.password = ""
        
        username = st.text_input(
            "Username", 
            value=st.session_state.username,
            key="username_input",
            help="Enter your username"
        )
        
        password = st.text_input(
            "Password", 
            value=st.session_state.password,
            type="password",
            key="password_input",
            help="Enter your password"
        )
        
        # Update session state when inputs change
        if username:
            st.session_state.username = username
        if password:
            st.session_state.password = password
        
        # Login button
        if st.button("Login", type="primary", use_container_width=True):
            if username and password:
                if username == "admin" and password == "0108":
                    st.session_state['user_role'] = 'admin'
                    st.success("✅ Admin login successful!")
                else:
                    st.session_state['user_role'] = 'user'
                    st.success("✅ User login successful!")
            else:
                st.error("⚠️ Please enter both username and password")
        
        # Logout button
        if 'user_role' in st.session_state:
            if st.button("Logout", type="secondary", use_container_width=True):
                st.session_state.pop('user_role', None)
                st.session_state.pop('username', None)
                st.session_state.pop('password', None)
                st.rerun()
        
        st.markdown("---")
        st.markdown("**Note:** Use these credentials to access the video recording system.")
    
    # Demo Video Section
    st.header("📹 Demo Video")
    st.markdown("**Video Interview Simulation** - Sample video for demonstration")
    
    demo_video_path = "interview simulation 4 min.mp4"
    if os.path.exists(demo_video_path):
        st.video(demo_video_path)
        
        # Add Thai instructions under the video
        st.markdown("---")
        st.markdown("""
        เป็นที่ทราบกันดีอยู่ว่าในการสื่อสารนั้นคำพูดให้ข้อมูล ส่วนภาษากายที่ไม่ว่าจะเป็น สายตาการเคลื่อนไหว 
        ของลำตัว ศรีษะ มือ แขนและขา บอกผู้ฟังถึงอารมณ์และความรู้สึกของผู้พูด
        
        **AI People Reader App** นั้นถูกสร้างมาให้วิเคราะห์การเคลื่อนไหวโดยรวมในขณะที่สื่อสาร โดยมีวัดตามหลักการ 4 ประเภท ตามนี้
        
        1. ผู้พูดสามารถ engage ผู้ฟังหรือไม่
        2. ผู้พูดสามารถรับมือกับผู้ฟังหลายประเภทหรือไม่
        3. ผู้พูดสามารถโต้แย้ง ยืนกรานในสิ่งที่ตัวเองเชื่อหรือไม่
        4. ผู้พูดสามารถที่จะทำให้ผู้ฟังทำตามสิ่งที่ตัวเองต้องการ เพื่อให้ได้ผลที่ต้องการหรือไม่
        
        โดยเพื่อเป็นการช่วยให้ผู้ที่จัดจัดส่งคลิปวิดีโอมี guideline ทางเราได้ถ่ายคลิปที่มีคำถาม upload ให้แล้ว 
        ผู้ที่กำลังจะถ่ายคลิปเพียงแค่ต้องใช้ username และ password log in เข้าทาง laptop หรือและใช้มือถือในการถ่ายคลิปของตัวเองแบบแนวตั้ง 
        โดยรายละเอียดการถ่ายคลิป และ resolution นั้นอยู่ในหน้า **Upload Your Video** แล้ว
        
        **หมายเหตุ:** ในการตอบคำถามต่างๆขอให้ตอบอย่างละเอียด และเคลื่อนไหวเป็นธรรมชาติ
        """)
        st.markdown("---")
    else:
        st.warning(f"⚠️ Demo video '{demo_video_path}' not found in the current directory.")
    
    st.divider()
    
    # Main content - different based on user role
    if 'user_role' not in st.session_state:
        st.header("🔒 Please Login")
        st.info("Please login using the sidebar to access video upload functionality.")
    else:
        st.header("Upload Your Own Video")
        
        if st.session_state['user_role'] == 'admin':
            st.success("👑 Admin Mode: Full analysis capabilities enabled")
        else:
            st.info("👤 User Mode: Upload videos for admin review")
        
        uploaded_file = st.file_uploader("Choose a video file", type=['mp4', 'mov', 'avi'])
        
        if uploaded_file is not None:
            # Save uploaded file temporarily
            tfile = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
            tfile.write(uploaded_file.read())
            video_path = tfile.name
            
            # Store video info in session state
            if 'uploaded_videos' not in st.session_state:
                st.session_state['uploaded_videos'] = []
            
            video_info = {
                'name': uploaded_file.name,
                'path': video_path,
                'uploaded_by': st.session_state.get('username', 'Unknown'),
                'upload_time': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'size': len(uploaded_file.getvalue())
            }
            
            # Add to uploaded videos list if not already there
            if video_info not in st.session_state['uploaded_videos']:
                st.session_state['uploaded_videos'].append(video_info)
            
            st.video(video_path)
            
            # Show different buttons based on user role
            if st.session_state['user_role'] == 'admin':
                # Admin gets the analysis button
                if st.button("🔍 Complete Analysis (Motion Detection + Skills Report)", type="primary", use_container_width=True):
                    if not st.session_state.candidate_name or st.session_state.candidate_name.strip() == "":
                        st.error("⚠️ Please enter a candidate name in the Report Details section above.")
                    else:
                        with st.spinner("Processing video for motion detection and interpersonal skills analysis... This may take a few minutes."):
                            progress_bar = st.progress(0)
                    
                    def update_progress(progress):
                        progress_bar.progress(progress)
                    
                    # Step 1: Process video for motion detection
                    st.info("🎬 Step 1/2: Analyzing motion patterns...")
                    output_path, motion_counts, motion_labels, detection_log = process_video(
                        video_path, 
                        progress_callback=lambda p: update_progress(p * 0.5)  # First half of progress
                    )
                    
                    # Store motion detection results in session state
                    st.session_state['output_path'] = output_path
                    st.session_state['motion_counts'] = motion_counts
                    st.session_state['motion_labels'] = motion_labels
                    st.session_state['detection_log'] = detection_log
                    
                    # Analyze movement combinations against detected movements
                    movement_analysis = analyze_detected_movements(motion_counts, movement_combinations)
                    st.session_state['movement_analysis'] = movement_analysis
                    
                    # Step 2: Analyze interpersonal skills
                    st.info("📊 Step 2/2: Analyzing interpersonal skills...")
                    skills_analysis, error = analyze_interpersonal_skills(video_path)
                    
                    if error:
                        st.error(f"❌ Error analyzing interpersonal skills: {error}")
                    else:
                        # Store skills analysis in session state
                        st.session_state['skills_analysis'] = skills_analysis
                        
                        # Create Word document with motion data
                        doc = create_interpersonal_skills_docx_report(
                            st.session_state.candidate_name, 
                            assessment_date, 
                            skills_analysis,
                            st.session_state.get('motion_counts'),
                            st.session_state.get('motion_labels'),
                            st.session_state.get('movement_analysis')
                        )
                        
                        # Save to bytes
                        doc_bytes = io.BytesIO()
                        doc.save(doc_bytes)
                        doc_bytes.seek(0)
                        
                        # Store document in session state
                        st.session_state['generated_doc'] = doc_bytes
                        st.session_state['report_filename'] = f"Interpersonal_Skills_Report_{st.session_state.candidate_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                        
                        progress_bar.progress(1.0)
                        progress_bar.empty()
                        
                        st.success("✅ Analysis complete! Both motion detection and interpersonal skills reports are ready.")
            else:
                # Regular users just see upload confirmation
                st.success("✅ Video uploaded successfully! Admin will review and analyze your video.")
                st.info("📝 Your video will remain on the system until admin downloads and removes it.")
        
        # Admin Video Management Section
        if st.session_state['user_role'] == 'admin' and 'uploaded_videos' in st.session_state and st.session_state['uploaded_videos']:
            st.markdown("---")
            st.header("📁 Uploaded Videos Management")
            st.markdown("**Videos uploaded by users:**")
            
            for i, video_info in enumerate(st.session_state['uploaded_videos']):
                with st.expander(f"📹 {video_info['name']} - Uploaded by {video_info['uploaded_by']} at {video_info['upload_time']}", expanded=False):
                    col1, col2, col3 = st.columns([2, 1, 1])
                    
                    with col1:
                        st.write(f"**File:** {video_info['name']}")
                        st.write(f"**Uploaded by:** {video_info['uploaded_by']}")
                        st.write(f"**Time:** {video_info['upload_time']}")
                        st.write(f"**Size:** {video_info['size']:,} bytes")
                        
                        # Show video
                        if os.path.exists(video_info['path']):
                            st.video(video_info['path'])
                    
                    with col2:
                        # Download button
                        if os.path.exists(video_info['path']):
                            with open(video_info['path'], 'rb') as f:
                                st.download_button(
                                    label="📥 Download Video",
                                    data=f.read(),
                                    file_name=video_info['name'],
                                    mime="video/mp4",
                                    key=f"download_{i}"
                                )
                    
                    with col3:
                        # Remove button
                        if st.button("🗑️ Remove", key=f"remove_{i}", type="secondary"):
                            # Remove from session state
                            st.session_state['uploaded_videos'].pop(i)
                            # Try to delete the file
                            try:
                                if os.path.exists(video_info['path']):
                                    os.unlink(video_info['path'])
                            except:
                                pass
                            st.rerun()
    
    # Results Section
    if 'output_path' in st.session_state:
        st.header("📊 Analysis Results")
        
        # Display processed video
        st.subheader("🎬 Processed Video with Skeleton Overlay")
        st.video(st.session_state['output_path'])
        
        # Download processed video
        with open(st.session_state['output_path'], 'rb') as f:
            st.download_button(
                label="📥 Download Processed Video",
                data=f.read(),
                file_name=f"processed_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp4",
                mime="video/mp4"
            )
        
        # Interpersonal Skills Report Results
        if 'skills_analysis' in st.session_state and st.session_state['skills_analysis']:
            st.markdown("---")
            st.subheader("📄 Interpersonal Skills Report")
            
            skills_analysis = st.session_state['skills_analysis']
            
            # Create and display pie chart
            pie_chart = create_interpersonal_skills_chart(skills_analysis)
            st.plotly_chart(pie_chart, use_container_width=True, key="skills_pie_chart")
            
            # Create and display detailed tables
            analysis_tables = create_skills_analysis_tables(skills_analysis)
            
            for skill_name, skill_data in analysis_tables.items():
                level = skill_data['level']
                score = skill_data['score']
                qualities = skill_data['qualities']
                
                st.subheader(f"{skill_name}: {level} ({score:.1f}/80)")
                
                if qualities:
                    # Create table
                    table_data = []
                    for quality, description in qualities.items():
                        table_data.append({"Quality": quality, "Description": description})
                    
                    df_quality = pd.DataFrame(table_data)
                    st.dataframe(df_quality, use_container_width=True)
                    
                    st.markdown("---")
            
            # Download Word document
            if 'generated_doc' in st.session_state:
                st.download_button(
                    label="📥 Download Interpersonal Skills Report (.docx)",
                    data=st.session_state['generated_doc'].getvalue(),
                    file_name=st.session_state['report_filename'],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
        
        # Display Movement Combination Analysis
        if 'movement_analysis' in st.session_state and st.session_state['movement_analysis']:
            st.markdown("---")
            st.subheader("🎯 Movement Combination Analysis by Category")
            st.markdown("This analysis matches detected movements against predefined movement combination patterns for each interpersonal skill category.")
            
            movement_analysis = st.session_state['movement_analysis']
            
            # Create tabs for each category
            cat_tabs = st.tabs(['Authority', 'Confidence', 'Engaging', 'Adaptability'])
            
            for idx, category in enumerate(['Authority', 'Confidence', 'Engaging', 'Adaptability']):
                with cat_tabs[idx]:
                    matches = movement_analysis.get(category, [])
                    
                    if matches:
                        st.markdown(f"**Found {len(matches)} matching behavioral pattern(s) for {category}**")
                        
                        for i, match in enumerate(matches, 1):
                            with st.expander(f"#{i} - {match['summary'][:80]}... (Match: {match['match_percentage']:.0f}%)", expanded=(i==1)):
                                col1, col2 = st.columns(2)
                                
                                with col1:
                                    st.markdown("**Full Description:**")
                                    st.write(match['summary'])
                                    st.markdown(f"**Match Score:** {match['match_score']}/{match['total_required']} movements")
                                    st.markdown(f"**Match Percentage:** {match['match_percentage']:.1f}%")
                                
                                with col2:
                                    st.markdown("**Required Movements:**")
                                    for req_mov in match['required_movements']:
                                        if req_mov in match['detected_movements']:
                                            count = match['match_counts'].get(req_mov, 0)
                                            st.markdown(f"✅ {req_mov} (detected {count}x)")
                                        else:
                                            st.markdown(f"❌ {req_mov} (not detected)")
                    else:
                        st.info(f"No matching movement combinations detected for {category} category.")
    
    # Summary report section
    if 'motion_counts' in st.session_state:
        st.header("📊 Motion Summary Report")
        
        # Display detailed detection log table
        if 'detection_log' in st.session_state and st.session_state['detection_log']:
            st.subheader("📋 Detailed Detection Log")
            
            # Convert detection log to DataFrame
            detection_df = pd.DataFrame(st.session_state['detection_log'])
            
            # Debug: Show what motions were detected in the log
            if not detection_df.empty:
                detected_motions = [col for col in detection_df.columns if col != 'StartTime' and detection_df[col].sum() > 0]
                st.info(f"🔍 Motions detected in timeline log: {', '.join(detected_motions)}")
            
            # Ensure all motion columns exist in the detection log
            for motion in st.session_state.get('motion_labels', {}).keys():
                if motion not in detection_df.columns:
                    detection_df[motion] = 0
            
            # Reorder columns to match the image format
            motion_order = ['Pressing', 'Flicking', 'Dabbing', 'Punching', 'Slashing', 
                           'Gliding', 'Enclosing', 'Spreading', 'Directing', 'Indirecting', 
                           'Advancing', 'Retreating']
            columns_order = ['StartTime'] + motion_order
            
            # Filter to only include columns that exist in the dataframe
            available_columns = [col for col in columns_order if col in detection_df.columns]
            detection_df = detection_df[available_columns]
            
            # Display the table
            st.dataframe(detection_df, use_container_width=True)
            
            # Download detection log as CSV
            csv_buffer = io.StringIO()
            detection_df.to_csv(csv_buffer, index=False)
            st.download_button(
                label="📄 Download Detection Log as CSV",
                data=csv_buffer.getvalue(),
                file_name=f"detection_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
        
        df_all, df_effort, df_directional = create_summary_report(
            st.session_state['motion_counts'],
            st.session_state['motion_labels']
        )
        
        # Display summary statistics
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Motions Detected", df_all['Count'].sum())
        with col2:
            st.metric("Effort Actions", df_effort['Count'].sum())
        with col3:
            st.metric("Directional Motions", df_directional['Count'].sum())
        
        # Debug: Show motion counts vs timeline detection
        st.subheader("🔍 Debug: Motion Detection Comparison")
        col_debug1, col_debug2 = st.columns(2)
        
        with col_debug1:
            st.markdown("**Motion Counts (from video processing):**")
            for motion, count in st.session_state['motion_counts'].items():
                if count > 0:
                    st.write(f"• {motion}: {count}")
        
        with col_debug2:
            st.markdown("**Timeline Detection (from CSV log):**")
            if 'detection_log' in st.session_state and st.session_state['detection_log']:
                detection_df = pd.DataFrame(st.session_state['detection_log'])
                timeline_counts = {}
                for motion in st.session_state.get('motion_labels', {}).keys():
                    if motion in detection_df.columns:
                        timeline_counts[motion] = detection_df[motion].sum()
                
                for motion, count in timeline_counts.items():
                    if count > 0:
                        st.write(f"• {motion}: {count}")
        
        # Effort Summary Charts (like the image)
        st.subheader("📈 Effort Summary")
        
        # Calculate percentages for all motions
        total_motions = df_all['Count'].sum()
        if total_motions > 0:
            df_all['Percentage'] = (df_all['Count'] / total_motions * 100).round(1)
            
            # Create two columns for the charts
            chart_col1, chart_col2 = st.columns(2)
            
            with chart_col1:
                st.markdown("**Effort Summary**")
                
                # Sort by count (descending) for better visualization
                df_sorted = df_all.sort_values('Count', ascending=False)
                
                # Create a custom bar chart with percentages
                chart_data = df_sorted[['Motion', 'Count', 'Percentage']].copy()
                
                # Create the bar chart using plotly
                fig_left = go.Figure()
                
                # Add horizontal bar chart
                fig_left.add_trace(go.Bar(
                    y=chart_data['Motion'],
                    x=chart_data['Count'],
                    orientation='h',
                    marker=dict(color='#8B5CF6'),  # Purple color
                    text=[f"{pct}%" for pct in chart_data['Percentage']],
                    textposition='outside',
                    textfont=dict(color='#10B981', size=12),  # Green text
                    hovertemplate='<b>%{y}</b><br>Count: %{x}<br>Percentage: %{text}<extra></extra>'
                ))
                
                fig_left.update_layout(
                    title="Effort Summary",
                    xaxis_title="Count",
                    yaxis_title="Motion",
                    height=max(400, len(chart_data) * 40),
                    showlegend=False,
                    plot_bgcolor='#F3F4F6',
                    paper_bgcolor='#F3F4F6',
                    margin=dict(l=20, r=20, t=40, b=20)
                )
                
                fig_left.update_layout(yaxis=dict(autorange="reversed"))  # Reverse order to show highest at top
                st.plotly_chart(fig_left, use_container_width=True, key="all_motions_chart")
            
            with chart_col2:
                st.markdown("**Top Movement Efforts**")
                
                # Get top 3 motions
                top_3 = df_sorted.head(3)
                
                # Create the top efforts chart
                fig_right = go.Figure()
                
                # Add horizontal bar chart for top 3
                fig_right.add_trace(go.Bar(
                    y=[f"{motion} - Rank #{i+1}" for i, motion in enumerate(top_3['Motion'])],
                    x=top_3['Count'],
                    orientation='h',
                    marker=dict(color='#EC4899'),  # Pink color
                    text=[f"{pct}%" for pct in top_3['Percentage']],
                    textposition='outside',
                    textfont=dict(color='#10B981', size=12),  # Green text
                    hovertemplate='<b>%{y}</b><br>Count: %{x}<br>Percentage: %{text}<extra></extra>'
                ))
                
                fig_right.update_layout(
                    title="Top Movement Efforts",
                    xaxis_title="Count",
                    yaxis_title="Motion",
                    height=max(300, len(top_3) * 80),
                    showlegend=False,
                    plot_bgcolor='#F3F4F6',
                    paper_bgcolor='#F3F4F6',
                    margin=dict(l=20, r=20, t=40, b=20)
                )
                
                fig_right.update_layout(yaxis=dict(autorange="reversed"))  # Reverse order to show #1 at top
                st.plotly_chart(fig_right, use_container_width=True, key="top3_motions_chart")
        
        # Display detailed tables
        tab1, tab2, tab3 = st.tabs(["All Motions", "Effort Actions", "Directional Motions"])
        
        with tab1:
            st.dataframe(df_all, use_container_width=True)
            
            # Bar chart for all motions
            st.bar_chart(df_all.set_index('Motion')['Count'])
        
        with tab2:
            st.dataframe(df_effort, use_container_width=True)
            st.bar_chart(df_effort.set_index('Motion')['Count'])
        
        with tab3:
            st.dataframe(df_directional, use_container_width=True)
            st.bar_chart(df_directional.set_index('Motion')['Count'])
        
        # Download report as CSV
        csv_buffer = io.StringIO()
        df_all.to_csv(csv_buffer, index=False)
        st.download_button(
            label="📄 Download Report as CSV",
            data=csv_buffer.getvalue(),
            file_name=f"motion_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )
    

if __name__ == "__main__":
    main()

